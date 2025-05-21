import os
import docx
import pandas as pd
import requests
from bs4 import BeautifulSoup

from langchain_ollama import OllamaEmbeddings, OllamaLLM
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.vectorstores import InMemoryVectorStore
from langchain_core.documents import Document
from unstructured.partition.pdf import partition_pdf
from unstructured.partition.utils.constants import PartitionStrategy

# Directories
data = './data/'
media = './media/'

os.makedirs(data, exist_ok=True)
os.makedirs(media, exist_ok=True)

# LLMs
llm = OllamaLLM(model="phi")
img_llm = OllamaLLM(model="llava")

# Embeddings and vector store
embeddings = OllamaEmbeddings(model="phi")
vector_store = InMemoryVectorStore(embeddings)

# Prompt template
template = """You are a helpful but concise assistant. Answer only using the provided context.
If the context is not relevant or does not answer the question, say "I don't know."
Keep your answers direct and under 3 sentences.
Question: {question}
Context: {context}
Answer:"""

# File parsing
def save_file(file):
    file_path = os.path.join(data, file.name)
    with open(file_path, "wb") as f:
        f.write(file.getbuffer())
    return file_path

def parse_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return parse_pdf(file_path)
    elif ext in [".jpg", ".jpeg", ".png"]:
        return parse_image(file_path)
    elif ext == ".txt":
        return open(file_path, "r", encoding="utf-8").read()
    elif ext == ".docx":
        return parse_docx(file_path)
    elif ext == ".csv":
        return parse_csv(file_path)
    else:
        return "[Unsupported file type]"

def parse_pdf(file_path):
    elements = partition_pdf(
        file_path,
        strategy=PartitionStrategy.HI_RES,
        extract_image_block_types=["Image", "Table"],
        extract_image_block_output_dir=media
    )
    text_elements = [e.text for e in elements if e.text and e.category not in ["Image", "Table"]]

    for f in os.listdir(media):
        if f.lower().endswith(".jpg"):
            try:
                image_path = os.path.join(media, f)
                image_text = text_from_image(image_path)
                text_elements.append(image_text)
            except Exception as e:
                text_elements.append(f"[Image error: {str(e)}]")

    return "\n".join(text_elements)

def parse_image(file_path):
    try:
        model = img_llm.bind(images=[file_path])
        return model.invoke("Describe this image in detail.")
    except Exception:
        return "[Image parsing error]"

def parse_docx(file_path):
    doc = docx.Document(file_path)
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

def parse_csv(file_path):
    try:
        df = pd.read_csv(file_path)
        summary = f"The dataset contains {df.shape[0]} rows and {df.shape[1]} columns.\n"
        return summary + df.head(5).to_markdown(index=False)
    except Exception as e:
        return f"[CSV parsing error: {e}]"

def parse_webpage(url):
    try:
        response = requests.get(url)
        soup = BeautifulSoup(response.text, 'html.parser')
        paragraphs = soup.find_all('p')
        return "\n".join(p.get_text() for p in paragraphs if p.get_text())
    except Exception as e:
        return f"[Web parsing error: {e}]"

# Vector operations
def split_text(text):
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    return splitter.split_text(text)

def store_chunks(chunks, source=None):
    docs = [
        Document(page_content=chunk, metadata={"source": source} if source else {})
        for chunk in chunks
    ]
    vector_store.add_documents(docs)

def retrieve_chunks(query):
    return vector_store.similarity_search(query, k=2)

def answer_question(question, docs):
    if not docs:
        return "I don't know."
    context = "\n\n".join([doc.page_content for doc in docs])
    sources = set(doc.metadata.get("source", "Unknown") for doc in docs)
    prompt = ChatPromptTemplate.from_template(template)
    chain = prompt | llm
    answer = chain.invoke({"question": question, "context": context})
    return answer + f"\n\nSources: {', '.join(sources)}"

def reset_vector_store():
    global vector_store
    vector_store = InMemoryVectorStore(embeddings)

# Reusable upload handler
def handle_multiple_files(file_list, parse_func):
    success_files = []
    for file in file_list:
        file_path = save_file(file)
        content = parse_func(file_path)
        if content.strip() and "Unsupported" not in content:
            chunks = split_text(content)
            store_chunks(chunks, source=file.name)
            success_files.append(file.name)
    return success_files
